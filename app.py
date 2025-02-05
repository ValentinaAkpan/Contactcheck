import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re
from thefuzz import fuzz

# List of provinces and personal names to exclude
EXCLUDE_NAMES = {
    "northwest territories", "saskatchewan", "new brunswick", "manitoba", "quebec",
    "new foundland", "nova scotia", "british columbia", "alberta", "ontario", "branches",
    "gary sandlac", "rick farrell", "brad cook", "kyle sayer", "steve fry"
}

# Function to normalize company names
def normalize_name(name):
    if not isinstance(name, str):
        return ""
    name = name.strip().lower()  # Remove spaces and convert to lowercase
    name = name.replace("&", "and")  # Replace '&' with 'and'
    name = re.sub(r'[^\w\s]', '', name)  # Remove punctuation
    return name

# Function to load all sheets in the Excel file
def load_excel(file):
    all_sheets = pd.read_excel(file, sheet_name=None, dtype=str)  # Load all sheets as a dictionary
    combined_data = pd.DataFrame()  # Placeholder for merged data

    # List of possible company name columns
    possible_names = {"Company", "Company Name", "Business Name", "Firm", "Member Name"}

    for sheet_name, df in all_sheets.items():
        df = df.fillna('')  # Replace NaN values with empty strings
        df = df.applymap(lambda x: x.strip() if isinstance(x, str) else x)  # Trim spaces

        found_columns = [col for col in df.columns if col.strip() in possible_names]

        if found_columns:
            combined_data = pd.concat([combined_data, df[found_columns]], ignore_index=True)

    # Normalize and return unique company names
    return {normalize_name(name) for col in combined_data.columns for name in combined_data[col].dropna()}

# Function to extract only company names from the Word document
def extract_company_names(doc):
    company_names = set()
    lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]  # Remove empty lines

    i = 0
    while i < len(lines):
        text = lines[i]  # Get first line (assumed company name)

        # Ignore lines that contain numbers, emails, or websites
        if re.search(r"\d", text) or "@" in text or "www." in text or ".com" in text or ".ca" in text:
            i += 1
            continue

        # Ignore common address indicators (e.g., "Street", "Avenue", "Drive", etc.)
        if re.search(r"\b(Street|St\.|Avenue|Ave\.|Drive|Dr\.|Road|Rd\.|Blvd|Highway|Hwy|Suite|PO Box|Postal)\b", text, re.IGNORECASE):
            i += 1
            continue

        # Normalize and exclude unwanted names
        normalized_text = normalize_name(text)
        if normalized_text not in EXCLUDE_NAMES:
            company_names.add(normalized_text)

        # Skip the next few lines (typically address, contact info)
        i += 4  

    return company_names

# Function to perform fuzzy matching (handles slight name variations)
def fuzzy_match(set1, set2, threshold=90):
    matched = set()
    for name1 in set1:
        for name2 in set2:
            if fuzz.ratio(name1, name2) >= threshold:
                matched.add(name1)
                break
    return matched

# Function to generate a Word document for updates
def generate_word_update(missing_in_word, extra_in_word):
    doc = Document()
    doc.add_heading("Membership Directory Updates", level=1)

    # Companies to Add
    if missing_in_word:
        doc.add_heading("Companies to Add:", level=2)
        for company in sorted(missing_in_word):
            doc.add_paragraph(company)

    # Companies to Remove
    if extra_in_word:
        doc.add_heading("Companies to Remove:", level=2)
        for company in sorted(extra_in_word):
            doc.add_paragraph(company)

    # Save to memory
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output

# Streamlit UI
st.title("Membership Directory Checker")
st.write("Upload an **Excel** file and an **existing Word document** to compare directory information.")

# File Uploads
excel_file = st.file_uploader("Upload Updated Excel File", type=["xlsx"])
word_file = st.file_uploader("Upload Existing Word Document", type=["docx"])

if excel_file and word_file:
    st.success("Files uploaded successfully!")

    # Load Excel data
    excel_companies = load_excel(excel_file)
    st.write(f"Total Companies in Excel (from all sheets): {len(excel_companies)}")

    # Load Word file and extract company names
    word_doc = Document(word_file)
    word_companies = extract_company_names(word_doc)
    st.write(f"Total Companies in Word Document: {len(word_companies)}")

    # Identify missing companies using exact and fuzzy matching
    missing_in_word = excel_companies - word_companies
    extra_in_word = word_companies - excel_companies

    # Apply fuzzy matching to reduce false positives
    fuzzy_matched = fuzzy_match(missing_in_word, word_companies)
    missing_in_word -= fuzzy_matched  # Remove fuzzy-matched names from missing list

    fuzzy_matched_extra = fuzzy_match(extra_in_word, excel_companies)
    extra_in_word -= fuzzy_matched_extra  # Remove fuzzy-matched names from extra list

    # Generate Word document if there are changes
    if missing_in_word or extra_in_word:
        word_update = generate_word_update(missing_in_word, extra_in_word)
        
        st.download_button(
            label="Download Word Update File",
            data=word_update,
            file_name="membership_updates.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Display changes
    st.write("### Missing Companies (Need to be Added)")
    if missing_in_word:
        st.write(list(missing_in_word))
    else:
        st.write("No new companies missing.")

    st.write("### Companies in Word but Not in Excel (Possible Removals)")
    if extra_in_word:
        st.write(list(extra_in_word))
    else:
        st.write("No companies need to be removed.")
